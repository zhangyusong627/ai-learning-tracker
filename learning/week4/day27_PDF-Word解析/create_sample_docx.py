from pathlib import Path

from docx import Document


output_path = Path(__file__).parent / "sample_contract.docx"

document = Document()

document.add_heading("软件服务合同", level=1)

document.add_paragraph("甲方：示例科技有限公司")
document.add_paragraph("乙方：测试软件有限公司")
document.add_paragraph("本合同用于约定双方的软件开发服务内容。")

document.add_heading("服务明细", level=2)

table = document.add_table(rows=1, cols=2)

header_cells = table.rows[0].cells
header_cells[0].text = "服务项目"
header_cells[1].text = "价格"

data = [
    ("系统开发", "100000元"),
    ("系统维护", "20000元"),
]

for service_name, price in data:
    cells = table.add_row().cells
    cells[0].text = service_name
    cells[1].text = price

document.add_paragraph("补充说明：以上价格均包含一年维护服务。")

document.save(output_path)

print(f"测试文档已生成：{output_path}")