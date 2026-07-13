from pathlib import Path

import fitz
from langchain_core.documents import Document as LangChainDocument
from docx import Document as WordDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_text_splitters import RecursiveCharacterTextSplitter


def iter_word_blocks(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def load_pdf(file_path: Path) -> list[LangChainDocument]:
    documents = []

    with fitz.open(file_path) as pdf:
        for page_index, page in enumerate(pdf):
            text = page.get_text().strip()

            if not text:
                print(f"警告：第 {page_index + 1} 页未提取到文字")
                continue

            documents.append(
                LangChainDocument(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_type": "pdf",
                        "page": page_index + 1,
                    },
                )
            )

    return documents


def load_word(file_path: Path) -> list[LangChainDocument]:
    documents = []
    word = WordDocument(file_path)

    for block_index, block in enumerate(
        iter_word_blocks(word),
        start=1,
    ):
        if isinstance(block, Paragraph):
            text = block.text.strip()

            if not text:
                continue

            documents.append(
                LangChainDocument(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_type": "docx",
                        "block_type": "paragraph",
                        "block_index": block_index,
                        "style": block.style.name,
                    },
                )
            )

        elif isinstance(block, Table):
            table_rows = []

            for row in block.rows:
                cell_values = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(cell_values):
                    table_rows.append(" | ".join(cell_values))

            if table_rows:
                documents.append(
                    LangChainDocument(
                        page_content="\n".join(table_rows),
                        metadata={
                            "source": file_path.name,
                            "file_type": "docx",
                            "block_type": "table",
                            "block_index": block_index,
                        },
                    )
                )

    return documents

DOCUMENT_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_word,
}


def load_document(file_path: Path) -> list[LangChainDocument]:
    suffix = file_path.suffix.lower()
    loader = DOCUMENT_LOADERS.get(suffix)

    if loader is None:
        print(f"跳过不支持的文件：{file_path.name}")
        return []

    return loader(file_path)


def load_directory(directory_path: Path):
    all_documents = []

    report = {
        "success_files": [],
        "failed_files": [],
        "skipped_files": [],
        "empty_files": [],
    }

    for file_path in sorted(directory_path.iterdir()):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()

        if suffix not in DOCUMENT_LOADERS:
            print(f"跳过不支持的文件：{file_path.name}")
            report["skipped_files"].append(file_path.name)
            continue

        try:
            file_documents = load_document(file_path)

            if not file_documents:
                print(f"未提取到内容：{file_path.name}")
                report["empty_files"].append(file_path.name)
                continue

            all_documents.extend(file_documents)

            report["success_files"].append({
                "file_name": file_path.name,
                "document_count": len(file_documents),
            })

            print(
                f"加载成功：{file_path.name}，"
                f"提取 {len(file_documents)} 条内容"
            )

        except Exception as error:
            report["failed_files"].append({
                "file_name": file_path.name,
                "error": str(error),
            })

            print(f"加载失败：{file_path.name}，原因：{error}")

    return all_documents, report

if __name__ == "__main__":
    documents_dir = Path(__file__).parent / "documents"

    all_documents, report = load_directory(documents_dir)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=[
            "\n\n",
            "\n",
            "。",
            "！",
            "？",
            "；",
            "，",
            " ",
            "",
        ],
    )

    chunks = text_splitter.split_documents(all_documents)

    print(f"切分前 Document 数量：{len(all_documents)}")
    print(f"切分后 Chunk 数量：{len(chunks)}")

    print("\n前 3 个 Chunk：")

    for index, chunk in enumerate(chunks[:3], start=1):
        print("-" * 50)
        print(f"Chunk {index}")
        print(f"字符数：{len(chunk.page_content)}")
        print(f"元数据：{chunk.metadata}")
        print(f"正文：{chunk.page_content[:200]}")

    print("\n批处理报告：")
    print(f"成功文件：{report['success_files']}")
    print(f"失败文件：{report['failed_files']}")
    print(f"跳过文件：{report['skipped_files']}")
    print(f"空内容文件：{report['empty_files']}")
