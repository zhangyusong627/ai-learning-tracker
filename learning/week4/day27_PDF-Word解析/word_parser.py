from pathlib import Path

from docx import Document


word_path = Path(__file__).parent / "sample_contract.docx"

document = Document(word_path)

print(f"普通段落数量：{len(document.paragraphs)}")
print(f"表格数量：{len(document.tables)}")
print("\n段落内容：")

for index, paragraph in enumerate(document.paragraphs, start=1):
    text = paragraph.text.strip()

    if not text:
        continue

    print(f"{index}. {text}")



print("\n表格内容：")

for table_index, table in enumerate(document.tables, start=1):
    print(f"表格 {table_index}：")

    for row_index, row in enumerate(table.rows, start=1):
        cell_values = [
            cell.text.strip()
            for cell in row.cells
        ]

        print(f"  第 {row_index} 行：{' | '.join(cell_values)}")

print("\nWord 底层元素顺序：")


for child in document.element.body.iterchildren():
    print(child.tag)