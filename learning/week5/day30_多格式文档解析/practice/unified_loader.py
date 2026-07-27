"""Day 30：PDF/Word 统一加载与解析质量报告。"""

from dataclasses import dataclass, field
from pathlib import Path
from zipfile import BadZipFile, ZipFile

import fitz
from docx import Document as WordDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document


@dataclass
class ParseResult:
    """一次文件解析的统一结果。"""

    file_name: str
    file_type: str
    status: str
    documents: list[Document] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    metrics: dict = field(default_factory=dict)


def detect_file_type(file_path: Path) -> str | None:
    """根据文件内容识别 PDF 或 DOCX，不只信任扩展名。"""
    with file_path.open("rb") as file:
        header = file.read(8)

    if header.startswith(b"%PDF-"):
        return "pdf"

    if header.startswith(b"PK"):
        try:
            with ZipFile(file_path) as archive:
                if "word/document.xml" in archive.namelist():
                    return "docx"
        except BadZipFile:
            return None

    return None


def common_metadata(file_path: Path, file_type: str) -> dict:
    return {
        "source": file_path.name,
        "file_type": file_type,
    }


def load_pdf(file_path: Path) -> ParseResult:
    documents = []
    empty_pages = []
    image_pages = []

    with fitz.open(file_path) as pdf:
        total_pages = len(pdf)

        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text().strip()

            if not text:
                empty_pages.append(page_number)
                if page.get_images(full=True):
                    image_pages.append(page_number)
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        **common_metadata(file_path, "pdf"),
                        "page": page_number,
                        "extraction_method": "text",
                    },
                )
            )

    warnings = []
    if empty_pages:
        warnings.append(f"未提取到文本的页码：{empty_pages}")

    if documents and empty_pages:
        status = "partial"
    elif documents:
        status = "success"
    elif image_pages:
        status = "ocr_required"
    else:
        status = "empty"

    return ParseResult(
        file_name=file_path.name,
        file_type="pdf",
        status=status,
        documents=documents,
        warnings=warnings,
        metrics={
            "total_pages": total_pages,
            "extracted_pages": len(documents),
            "empty_pages": empty_pages,
            "image_pages": image_pages,
            "total_characters": sum(len(doc.page_content) for doc in documents),
        },
    )


def iter_word_blocks(document: WordDocument):
    """按 Word XML 中的真实顺序返回段落和表格。"""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def load_docx(file_path: Path) -> ParseResult:
    word = WordDocument(file_path)
    documents = []
    empty_blocks = 0

    for block_index, block in enumerate(iter_word_blocks(word), start=1):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            block_type = "paragraph"
            extra_metadata = {"style": block.style.name}
        else:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in block.rows
            ]
            text = "\n".join(row for row in rows if row.strip(" |"))
            block_type = "table"
            extra_metadata = {}

        if not text:
            empty_blocks += 1
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    **common_metadata(file_path, "docx"),
                    "block_type": block_type,
                    "block_index": block_index,
                    **extra_metadata,
                },
            )
        )

    return ParseResult(
        file_name=file_path.name,
        file_type="docx",
        status="success" if documents else "empty",
        documents=documents,
        metrics={
            "extracted_blocks": len(documents),
            "empty_blocks": empty_blocks,
            "total_characters": sum(len(doc.page_content) for doc in documents),
        },
    )


LOADERS = {
    "pdf": load_pdf,
    "docx": load_docx,
}


def load_document(file_path: Path) -> ParseResult:
    """识别类型、校验后缀并分发到对应解析器。"""
    detected_type = detect_file_type(file_path)

    if detected_type is None:
        return ParseResult(
            file_name=file_path.name,
            file_type="unknown",
            status="skipped",
            warnings=["不支持或无法识别真实文件类型"],
        )

    result = LOADERS[detected_type](file_path)
    expected_suffix = f".{detected_type}"

    if file_path.suffix.lower() != expected_suffix:
        result.warnings.append(
            f"扩展名 {file_path.suffix or '(无)'} 与真实类型 {detected_type} 不一致"
        )

    return result


def print_result(result: ParseResult) -> None:
    print("=" * 60)
    print(f"文件：{result.file_name}")
    print(f"真实类型：{result.file_type}")
    print(f"状态：{result.status}")
    print(f"Document 数量：{len(result.documents)}")
    print(f"指标：{result.metrics}")
    print(f"警告：{result.warnings or '无'}")


def main() -> None:
    learning_dir = Path(__file__).resolve().parents[3]
    fixture_dir = learning_dir / "week4" / "day27_PDF-Word解析" / "documents"

    for file_path in sorted(fixture_dir.iterdir()):
        if file_path.is_file():
            print_result(load_document(file_path))


if __name__ == "__main__":
    main()
